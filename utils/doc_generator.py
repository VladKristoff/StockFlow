# utils/doc_generator.py
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
import os


def generate_invoice_document(doc_data, items, doc_type):
    """
    Создать Word документ накладной
    doc_type: "incoming" или "expenses"
    """
    # Создаем документ
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(f"{'ПРИХОДНАЯ НАКЛАДНАЯ' if doc_type == 'incoming' else 'РАСХОДНАЯ НАКЛАДНАЯ'}")
    title_run.font.size = Pt(16)
    title_run.font.bold = True

    doc.add_paragraph()

    # Шапка документа
    table_header = doc.add_table(rows=4, cols=2)
    table_header.style = 'Table Grid'

    # Заполняем шапку
    header_data = [
        ("Номер документа:", doc_data['doc_number']),
        ("Дата документа:", doc_data['doc_date']),
        ("Контрагент:", doc_data['contractor_name']),
        ("Ответственный:", doc_data['employee_name'])
    ]

    for i, (label, value) in enumerate(header_data):
        cell_label = table_header.cell(i, 0)
        cell_value = table_header.cell(i, 1)
        cell_label.text = label
        cell_value.text = str(value)
        cell_label.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # Таблица товаров
    doc.add_paragraph("Товары:")

    table_items = doc.add_table(rows=1, cols=5)
    table_items.style = 'Table Grid'

    # Заголовки таблицы
    headers = ['№', 'Наименование товара', 'Количество', 'Цена', 'Сумма']
    for i, header in enumerate(headers):
        cell = table_items.cell(0, i)
        cell.text = header
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Заполняем товарами
    total_sum = 0
    for idx, item in enumerate(items, 1):
        row = table_items.add_row()
        row.cells[0].text = str(idx)
        row.cells[1].text = item['product_name']
        row.cells[2].text = str(item['quantity'])
        row.cells[3].text = f"{item['price']:.2f} руб."
        total = item['quantity'] * item['price']
        row.cells[4].text = f"{total:.2f} руб."
        total_sum += total

        # Выравнивание
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Итоговая строка
    row_total = table_items.add_row()
    row_total.cells[0].text = ""
    row_total.cells[1].text = ""
    row_total.cells[2].text = ""
    row_total.cells[3].text = "ИТОГО:"
    row_total.cells[4].text = f"{total_sum:.2f} руб."
    for cell in row_total.cells:
        if cell.text:
            cell.paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph()

    # Подписи
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.add_run("______________________").font.bold = True
    doc.add_paragraph("(подпись, печать)")

    # Сохраняем документ
    filename = f"{doc_data['doc_number']}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    filepath = os.path.join("documents", filename)

    # Создаем папку если её нет
    os.makedirs("documents", exist_ok=True)

    doc.save(filepath)
    return filepath

def create_incoming_invoice(doc_id, doc_data, items):
    """Создать приходную накладную"""
    return generate_invoice_document(doc_data, items, "incoming")


def create_expense_invoice(doc_id, doc_data, items):
    """Создать расходную накладную"""
    return generate_invoice_document(doc_data, items, "expenses")